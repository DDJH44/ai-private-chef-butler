"""
Fill the AGI全栈开发大作业报告 template with project content.
Uses python-docx for paragraph-level text replacement and content insertion.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import copy

TEMPLATE = r"C:\Users\Lenovo\OneDrive\文档\xwechat_files\wxid_oflhucmainqz22_3c10\msg\file\2026-05\XXX项目小组-AGI全栈开发大作业报告（模板本） (2).docx"
OUTPUT = r"D:\AGI\AI Private Chef Butler\AI私人厨师-AGI全栈开发大作业报告.docx"

# ── Open template ──
doc = Document(TEMPLATE)

def find_para(doc, text_snippet, contains=True):
    """Find first paragraph containing text_snippet."""
    for i, p in enumerate(doc.paragraphs):
        if contains and text_snippet in p.text:
            return i, p
        elif not contains and p.text.strip() == text_snippet:
            return i, p
    return None, None

def replace_para_text(p, new_text, bold=False, size=None):
    """Replace all runs in a paragraph with new text, preserving first run's formatting."""
    if not p.runs:
        return
    # Keep first run's formatting as base
    r0 = p.runs[0]
    fmt = r0._element.find(qn('w:rPr'))
    # Clear all runs
    for r in p.runs:
        r.text = ""
    # Set text in first run
    r0.text = new_text
    if bold:
        r0.bold = True
    if size:
        r0.font.size = size

def set_para_text(p, text):
    """Completely replace all text in a paragraph."""
    # Clear existing runs
    for r in p.runs:
        r._element.getparent().remove(r._element)
    # Add new run
    if p.runs:
        p.runs[0].text = text
    else:
        # Find or create rPr from pPr
        pPr = p._element.find(qn('w:pPr'))
        rPr = None
        if pPr is not None:
            rPr = pPr.find(qn('w:rPr'))
        run = p.add_run(text)
        if rPr is not None:
            # Copy formatting
            pass

def find_and_replace(doc, search, replacement):
    """Replace text in the first paragraph containing search."""
    idx, para = find_para(doc, search)
    if para and para.runs:
        full_text = para.text
        new_text = full_text.replace(search, replacement)
        # Replace in runs
        remaining = replacement
        for r in para.runs:
            if search in r.text:
                r.text = r.text.replace(search, replacement)
                break
        # If replacement not fully placed, put in first run
        if search not in "".join(r.text for r in para.runs):
            pass  # Already done
    return para

def find_para_index(doc, text_snippet):
    """Return index of first paragraph containing the text."""
    for i, p in enumerate(doc.paragraphs):
        if text_snippet in p.text:
            return i
    return None


# ============================================================
# STEP 1: Replace cover page info
# ============================================================
# Project title: "XXX 项目架构设计及全栈开发" → "AI私人厨师 项目架构设计及全栈开发"
find_and_replace(doc, "XXX 项目架构设计及全栈开发", "AI私人厨师 项目架构设计及全栈开发")

# Semester: leave as-is (already 2024-2025 第一学期)

# Completion time: replace "完成时间 年 月 日"
find_and_replace(doc, "完成时间 年 月 日", "完成时间 2026年5月15日")

# ============================================================
# STEP 2: Fill Section 1 — 项目计划
# ============================================================

# 1.1 项目简介
idx, _ = find_para(doc, "包括你项目的名称，所属行业，并描述你App的主要功能和内容。")
if idx:
    p = doc.paragraphs[idx]
    for r in p.runs:
        r._element.getparent().remove(r._element)
    run = p.add_run(
        "AI私人厨师（AI Private Chef Butler）是一款面向家庭烹饪场景的智能饮食助手应用，"
        "所属行业为智慧生活/餐饮科技。用户可通过自然语言对话、拍照识别食材等方式与AI交互，"
        "获取个性化的菜谱推荐、营养分析、一周膳食规划和购物清单。"
        "系统基于大语言模型（豆包Seed系列）与LangGraph多智能体框架，"
        "能够根据用户冰箱库存、饮食偏好（过敏源、口味、饮食类型）和家庭成员结构，"
        "智能推荐合适的菜品，并自动搜索菜品实物照片和B站视频教程。"
        "前端采用Next.js + Tailwind CSS构建响应式交互界面，"
        "后端使用FastAPI提供高性能API服务，"
        "数据存储采用SQLite轻量数据库，支持Docker容器化部署。"
    )

# 1.2 项目背景
idx, _ = find_para(doc, "对系统的社会、经济、业务背景进行简单介绍。")
if idx:
    p = doc.paragraphs[idx]
    for r in p.runs:
        r._element.getparent().remove(r._element)
    run = p.add_run("")

# Find the "例如：" and "政策支持" paragraph, then insert background content
# We'll replace the "政策支持" paragraph and add content after it
idx_policy, _ = find_para(doc, "政策支持")
if idx_policy:
    p = doc.paragraphs[idx_policy]
    p.text = ""
    for r in p.runs:
        r._element.getparent().remove(r._element)
    run = p.add_run(
        "政策支持：国务院《"健康中国2030"规划纲要》明确提出推广健康生活方式，"
        "国家卫健委发布的《中国居民膳食指南（2022）》为科学饮食提供了权威标准。"
        "教育部等多部门推动劳动教育与烹饪技能培养，为智能饮食辅助工具创造了政策空间。"
    )

idx_industry, _ = find_para(doc, "行业现状")
if idx_industry:
    p = doc.paragraphs[idx_industry]
    p.text = ""
    for r in p.runs:
        r._element.getparent().remove(r._element)
    run = p.add_run(
        "行业现状：中国在线餐饮市场2025年规模已超5万亿元，但家庭烹饪场景的数字化渗透率仍不足15%。"
        "现有菜谱App（如下厨房、豆果美食）以UGC内容为主，缺乏个性化智能推荐能力。"
        "AI大模型技术的成熟使得"千人千面"的个性化饮食指导成为可能。"
        "参考文献：中国营养学会《中国居民膳食指南（2022）》；艾瑞咨询《2025中国在线餐饮行业研究报告》"
    )

# Find and remove the "等等" paragraph
idx_etc, _ = find_para(doc, "等等")
if idx_etc:
    p = doc.paragraphs[idx_etc]
    for r in p.runs:
        r._element.getparent().remove(r._element)

# 1.3.1 痛点问题/核心业务
idx_pain, _ = find_para(doc, "描述项目解决的痛点问题")
if idx_pain:
    p = doc.paragraphs[idx_pain]
    for r in p.runs:
        r._element.getparent().remove(r._element)
    run = p.add_run(
        "痛点一：不知道吃什么——家庭烹饪最大的痛点是"选择困难"，面对冰箱里的食材不知道能做什么菜。"
        "痛点二：不会做——很多用户想做某道菜但缺乏详细步骤指导，视频教程搜索耗时。"
        "痛点三：饮食不科学——缺乏营养知识导致膳食结构不均衡，无法针对过敏、减脂等特殊需求调整。"
        "痛点四：食材浪费——买回来的食材因不知道如何搭配而过期浪费。"
    )

idx_biz, _ = find_para(doc, "描述项目主要的业务逻辑等")
if idx_biz:
    p = doc.paragraphs[idx_biz]
    for r in p.runs:
        r._element.getparent().remove(r._element)
    run = p.add_run(
        "核心业务逻辑：用户上传食材照片或通过文字描述需求 → AI多模态识别食材/场景 → "
        "结合用户偏好（过敏源、口味、饮食类型）和冰箱库存 → 大语言模型生成个性化菜谱推荐 → "
        "自动搜索真实菜品照片和B站教学视频 → 结构化输出菜谱（食材、调料、步骤、评分）→ "
        "用户可选择保存菜谱、生成购物清单、纳入膳食计划。"
        "系统同时支持熟食营养分析：用户拍摄已完成的饭菜，AI识别后估算热量和营养素，给出健康评价。"
    )

# 1.3.2 产品结构图
idx_prod, _ = find_para(doc, "描述系统主要产品结构")
if idx_prod:
    p = doc.paragraphs[idx_prod]
    for r in p.runs:
        r._element.getparent().remove(r._element)
    run = p.add_run(
        "产品结构分为六大功能模块：\n"
        "1. AI智能对话 — 核心交互入口，支持文字+图片多模态输入\n"
        "   ├── 食材识别 → 菜谱推荐\n"
        "   ├── 熟食分析 → 营养评估\n"
        "   └── 情境感知 → 针对性推荐\n"
        "2. 菜谱管理 — 菜谱的保存、搜索、浏览、导出\n"
        "3. 冰箱库存 — 食材的增删改查、保质期追踪\n"
        "4. 膳食规划 — AI生成一周三餐计划、营养汇总\n"
        "5. 购物清单 — 基于菜谱自动生成采购清单、勾选管理\n"
        "6. 个人中心 — 饮食偏好设置、营养记录、烹饪历史、飞书集成"
    )

# 1.3.3 主要业务流程描述
idx_flow, _ = find_para(doc, "对系统功能描述中出现的主要业务流程进行描述。")
if idx_flow:
    p = doc.paragraphs[idx_flow]
    for r in p.runs:
        r._element.getparent().remove(r._element)
    run = p.add_run(
        "核心业务流程（食材识别→菜谱推荐）：\n"
        "1. 用户通过对话输入框发送消息，可选附带食材照片\n"
        "2. 前端将图片上传至阿里云OSS，获取图片URL\n"
        "3. 前端调用 /api/v1/chat/stream 接口，传入消息、图片URL、用户偏好、库存数据\n"
        "4. 后端构建完整上下文（系统提示词 + 用户偏好 + 冰箱库存 + 用户消息）\n"
        "5. LangGraph Agent 调用大模型进行多模态理解\n"
        "   a. 首先判断图片类型：熟食→营养分析模式；生食材→菜谱推荐模式\n"
        "   b. 食材模式下，模型先基于知识推荐菜谱\n"
        "   c. 为每个菜谱调用 recipe_search 工具搜索实物照片（Pexels/Unsplash API）\n"
        "   d. 视觉模型验证照片与菜品的匹配度\n"
        "   e. 调用 bilibili_search 工具搜索B站教学视频\n"
        "6. 模型按固定格式输出结构化菜谱（含SAVE_RECIPE标记块）\n"
        "7. 前端流式接收响应，实时渲染Markdown内容\n"
        "8. 前端解析SAVE_RECIPE标记，弹出批量保存确认面板\n"
        "9. 用户确认后调用 /api/v1/recipes/batch-create 保存菜谱"
    )

# 1.3.4 信息结构图
idx_info, _ = find_para(doc, "描述系统主要数据信息")
if idx_info:
    p = doc.paragraphs[idx_info]
    for r in p.runs:
        r._element.getparent().remove(r._element)
    run = p.add_run(
        "系统核心数据实体及关系：\n"
        "├── 用户（User）：id, 用户名, 邮箱, 密码哈希, 创建时间\n"
        "├── 菜谱（Recipe）：id, 用户ID, 标题, 正文, 图片URL, 难度, 烹饪时间, 食材列表, 调料列表, 标签, 评分, 来源URL, 视频URL\n"
        "├── 购物清单（ShoppingList）：id, 用户ID, 关联菜谱ID列表, 商品项列表（名称、数量、单位、是否勾选）, 状态\n"
        "├── 饮食偏好（Preference）：用户ID, 过敏源列表, 饮食类型, 口味偏好（辣/咸/甜/油程度）, 家庭成员\n"
        "├── 食材库存（Ingredient）：id, 用户ID, 名称, 分类, 数量, 单位, 购买日期, 保质期, 状态\n"
        "├── 营养记录（NutritionRecord）：id, 用户ID, 日期, 餐次类型, 食物名称, 热量, 蛋白质, 碳水, 脂肪, 纤维, 钠\n"
        "├── 烹饪历史（CookHistory）：id, 用户ID, 菜谱ID, 烹饪日期, 评分, 备注\n"
        "└── 飞书设置（FeishuSettings）：用户ID, 飞书用户标识, Webhook URL, 通知偏好\n"
        "核心关系：一个用户拥有多个菜谱、多条购物清单、多项营养记录；偏好和飞书设置与用户一一对应。"
    )

# ============================================================
# STEP 3: Fill Section 2 — 前端设计与开发
# ============================================================

# 2.1.1 模块命名与结构
idx_mod, _ = find_para(doc, "2.1.1 模块命名与结构")
if idx_mod:
    # Find next paragraph after heading (the placeholder content)
    for i in range(idx_mod + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip():
            p = doc.paragraphs[i]
            for r in p.runs:
                r._element.getparent().remove(r._element)
            run = p.add_run(
                "前端项目基于 Next.js 14 App Router 架构，模块组织结构如下：\n\n"
                "frontend/\n"
                "├── app/                    # 页面路由（Next.js App Router）\n"
                "│   ├── page.tsx            # 首页：AI对话交互\n"
                "│   ├── layout.tsx           # 根布局：全局样式+导航\n"
                "│   ├── recipes/             # 菜谱管理页\n"
                "│   ├── meal-plan/           # 膳食规划页\n"
                "│   ├── shopping-list/       # 购物清单页\n"
                "│   ├── fridge/              # 冰箱库存页\n"
                "│   ├── nutrition/           # 营养记录页\n"
                "│   ├── preferences/         # 饮食偏好设置页\n"
                "│   ├── profile/             # 个人中心页\n"
                "│   ├── history/             # 烹饪历史页\n"
                "│   ├── login/               # 登录页\n"
                "│   └── register/            # 注册页\n"
                "├── components/              # 可复用React组件\n"
                "├── lib/                     # 工具库（API调用、状态管理）\n"
                "│   ├── api.ts               # HTTP API封装\n"
                "│   └── chatStore.ts          # 聊天状态全局管理\n"
                "├── types/                   # TypeScript类型定义\n"
                "└── public/                  # 静态资源"
            )
            break

# 2.1.2 页面命名与结构
idx_page, _ = find_para(doc, "2.1.2 页面命名与结构")
if idx_page:
    for i in range(idx_page + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip():
            p = doc.paragraphs[i]
            for r in p.runs:
                r._element.getparent().remove(r._element)
            run = p.add_run(
                "页面路由与功能对应表：\n\n"
                "| 路由 | 页面名称 | 功能描述 |\n"
                "|------|----------|----------|\n"
                "| / | 首页（AI对话） | 核心交互页，多模态AI对话，食材识别与菜谱推荐 |\n"
                "| /recipes | 我的菜谱 | 菜谱的搜索、浏览、详情查看、删除管理 |\n"
                "| /meal-plan | 膳食规划 | AI生成一周早/午/晚餐计划，支持部分重新生成 |\n"
                "| /shopping-list | 购物清单 | 查看和管理基于菜谱生成的采购清单 |\n"
                "| /fridge | 冰箱管理 | 食材库存的增删改查、保质期管理 |\n"
                "| /nutrition | 营养日记 | 拍照分析饮食营养、查看每日营养摄入汇总 |\n"
                "| /preferences | 偏好设置 | 过敏源、饮食类型、口味、家庭成员信息管理 |\n"
                "| /profile | 个人中心 | 用户信息、飞书集成设置 |\n"
                "| /history | 烹饪历史 | 查看历史烹饪记录 |\n"
                "| /login | 登录 | 用户认证登录 |\n"
                "| /register | 注册 | 新用户注册 |"
            )
            break

# 2.2.1 页面：AI对话首页
# Find the second occurrence of "2.2.1" section or the placeholder after it
idx_page_xxx, _ = find_para(doc, "1) HTML页面模板DOM结构")
if idx_page_xxx:
    p = doc.paragraphs[idx_page_xxx]
    for r in p.runs:
        r._element.getparent().remove(r._element)
    # We need to insert structured content here
    # Since we can't easily add new paragraphs in place, replace with combined content
    pass

print("Phase 1 complete — basic text replacements done.")
print("Saving intermediate result...")

# Save progress
doc.save(OUTPUT)
print(f"Saved to: {OUTPUT}")
